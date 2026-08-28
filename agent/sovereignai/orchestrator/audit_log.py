"""
audit_log.py — Append-only JSONL + SQLite audit trail.

Every agent turn writes a complete record. Nothing is ever redacted.
This is the compliance trail — kept entirely local, never transmitted.
"""
from __future__ import annotations

import json
import os
import sqlite3
import time
import uuid
from pathlib import Path
from typing import Any

from sovereignai.config import get_config


def _db_path() -> Path:
    return get_config().audit_path / "audit.db"


def _jsonl_path() -> Path:
    return get_config().audit_path / "audit.jsonl"


def _ensure_db(conn: sqlite3.Connection) -> None:
    conn.execute("""
        CREATE TABLE IF NOT EXISTS audit_turns (
            id              TEXT PRIMARY KEY,
            session_id      TEXT NOT NULL,
            timestamp       REAL NOT NULL,
            user_text       TEXT,
            category        TEXT,
            model_name      TEXT,
            confidence      REAL,
            tool_calls_json TEXT,
            response_text   TEXT,
            token_count     INTEGER,
            duration_s      REAL,
            os_user         TEXT
        )
    """)
    conn.execute("""
        CREATE INDEX IF NOT EXISTS idx_session ON audit_turns(session_id)
    """)
    conn.commit()


class AuditLog:
    """Singleton-style audit logger. Use module-level `record()` instead."""

    def __init__(self) -> None:
        cfg = get_config()
        cfg.audit_path.mkdir(parents=True, exist_ok=True)
        self._db = sqlite3.connect(str(_db_path()), check_same_thread=False)
        _ensure_db(self._db)
        self._jsonl = open(_jsonl_path(), "a", encoding="utf-8")

    def record(
        self,
        session_id: str,
        user_text: str,
        category: str,
        model_name: str,
        confidence: float,
        response_text: str,
        tool_calls: list[dict[str, Any]],
        token_count: int = 0,
        duration_s: float = 0.0,
    ) -> str:
        record_id = str(uuid.uuid4())
        ts = time.time()
        os_user = os.environ.get("USER", os.environ.get("USERNAME", "unknown"))

        row = {
            "id": record_id,
            "session_id": session_id,
            "timestamp": ts,
            "user_text": user_text,
            "category": category,
            "model_name": model_name,
            "confidence": confidence,
            "tool_calls": tool_calls,
            "response_text": response_text,
            "token_count": token_count,
            "duration_s": duration_s,
            "os_user": os_user,
        }

        # JSONL — always written first, survives crashes
        self._jsonl.write(json.dumps(row, ensure_ascii=False) + "\n")
        self._jsonl.flush()

        # SQLite mirror
        self._db.execute("""
            INSERT INTO audit_turns VALUES (?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            record_id, session_id, ts, user_text,
            category, model_name, confidence,
            json.dumps(tool_calls), response_text,
            token_count, duration_s, os_user,
        ))
        self._db.commit()

        return record_id

    def get_session(self, session_id: str) -> list[dict[str, Any]]:
        cur = self._db.execute(
            "SELECT * FROM audit_turns WHERE session_id=? ORDER BY timestamp",
            (session_id,)
        )
        cols = [d[0] for d in cur.description]
        rows = []
        for row in cur.fetchall():
            d = dict(zip(cols, row))
            if d.get("tool_calls_json"):
                d["tool_calls"] = json.loads(d["tool_calls_json"])
                del d["tool_calls_json"]
            rows.append(d)
        return rows

    def list_sessions(self) -> list[dict[str, Any]]:
        cur = self._db.execute("""
            SELECT session_id,
                   MIN(timestamp) as started_at,
                   MAX(timestamp) as last_at,
                   COUNT(*) as turns
            FROM audit_turns
            GROUP BY session_id
            ORDER BY started_at DESC
        """)
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, row)) for row in cur.fetchall()]

    def close(self) -> None:
        self._jsonl.close()
        self._db.close()


# Module-level singleton
_audit: AuditLog | None = None


def get_audit() -> AuditLog:
    global _audit
    if _audit is None:
        _audit = AuditLog()
    return _audit


def record(*args: Any, **kwargs: Any) -> str:
    return get_audit().record(*args, **kwargs)


def export_session(
    session_id: str | None = None,
    fmt: str = "docx",
    output: str | None = None,
) -> Path:
    """Export a session audit record to DOCX, JSON, or plain text."""
    audit = get_audit()
    sessions = audit.list_sessions()
    if not sessions:
        raise ValueError("No sessions found in audit log.")

    if session_id is None:
        session_id = sessions[0]["session_id"]

    turns = audit.get_session(session_id)
    if not turns:
        raise ValueError(f"Session {session_id} not found.")

    cfg = get_config()
    out_dir = cfg.audit_path

    if fmt == "json":
        out_path = output or str(out_dir / f"audit_{session_id[:8]}.json")
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(turns, f, indent=2, ensure_ascii=False)

    elif fmt == "text":
        out_path = output or str(out_dir / f"audit_{session_id[:8]}.txt")
        with open(out_path, "w", encoding="utf-8") as f:
            for t in turns:
                f.write(f"[{t['timestamp']}] {t['os_user']} → {t['category']} ({t['model_name']})\n")
                f.write(f"USER: {t['user_text']}\n")
                f.write(f"RESPONSE: {t['response_text']}\n\n")

    elif fmt == "docx":
        out_path = output or str(out_dir / f"audit_{session_id[:8]}.docx")
        try:
            from docx import Document
            from docx.shared import Pt
            doc = Document()
            doc.add_heading("SovereignAI Audit Report", 0)
            doc.add_paragraph(f"Session ID: {session_id}")
            doc.add_paragraph(f"Turns: {len(turns)}")
            doc.add_paragraph("")
            for i, t in enumerate(turns, 1):
                import datetime
                ts_str = datetime.datetime.fromtimestamp(t["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
                doc.add_heading(f"Turn {i} — {ts_str}", level=2)
                doc.add_paragraph(f"User: {t['user_text']}")
                doc.add_paragraph(f"Category: {t['category']}  |  Model: {t['model_name']}  |  Confidence: {t.get('confidence', 0):.2f}")
                if t.get("tool_calls"):
                    doc.add_paragraph(f"Tool calls: {len(t['tool_calls'])}")
                doc.add_paragraph(f"Response:\n{t['response_text']}")
            doc.save(out_path)
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    else:
        raise ValueError(f"Unknown format: {fmt}. Use: docx | json | text")

    return Path(out_path)

